"""
document_processor.py — DOCX read/write with structure preservation.

Extracts text from paragraphs and table cells, runs PII detection per
text chunk, then applies replacements at the python-docx Run level so
formatting (bold, italic, font size, colour) is preserved as much as
possible.
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def load_document(path: str | Path) -> Document:
    """Load a DOCX file and return a python-docx Document object."""
    return docx.Document(str(path))


def iter_paragraphs(doc: Document):
    """
    Yield all paragraphs in the document, including those inside tables.
    Yields tuples of (paragraph, context_label) where context_label is
    e.g. 'body', 'table[0][0]'.
    """
    # Body paragraphs
    for para in doc.paragraphs:
        yield para, "body"

    # Table paragraphs
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    yield para, f"table[{t_idx}][{r_idx},{c_idx}]"


def get_paragraph_text(para: Paragraph) -> str:
    """Return the full text of a paragraph."""
    return para.text


def apply_replacements_to_paragraph(
    para: Paragraph,
    replacements: List[Tuple[str, str]],
) -> None:
    """
    Apply a list of (original, replacement) pairs to a paragraph.

    Strategy: merge all runs into a single run text, apply string
    replacements, then re-set the text back to the first run and clear
    the rest.  This is a pragmatic approach that preserves paragraph-level
    formatting but may lose inter-run italic/bold boundaries inside a
    single redacted token.
    """
    if not replacements:
        return

    # Collect full text and per-run data
    full_text = "".join(run.text for run in para.runs)
    if not full_text.strip():
        return

    modified = full_text
    for original, replacement in replacements:
        if original in modified:
            modified = modified.replace(original, replacement)

    if modified == full_text:
        return  # nothing changed

    # Write back: put modified text in run[0], clear the rest
    if para.runs:
        para.runs[0].text = modified
        for run in para.runs[1:]:
            run.text = ""
