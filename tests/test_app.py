"""
tests/test_app.py — Integration tests for FastAPI cloud web service.

Tests endpoints:
  - GET  /
  - GET  /health
  - POST /redact (synthetic mini DOCX)
"""

import io
import docx
import pytest
from fastapi.testclient import TestClient

from app import app


@pytest.fixture
def client():
    """Create a FastAPI TestClient."""
    return TestClient(app)


def _create_synthetic_docx(paragraphs: list[str]) -> io.BytesIO:
    """Create an in-memory DOCX file with given paragraph strings."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


class TestAppEndpoints:

    def test_get_index(self, client):
        """GET / returns HTML landing page."""
        response = client.get("/")
        assert response.status_code == 200
        assert "text/html" in response.headers.get("content-type", "")
        assert "PII Redaction Tool" in response.text
        assert "dropzone" in response.text

    def test_get_health(self, client):
        """GET /health returns JSON status ok."""
        response = client.get("/health")
        assert response.status_code == 200
        data = response.json()
        assert data.get("status") == "ok"
        assert data.get("service") == "pii-redactor"

    def test_redact_rejects_non_docx(self, client):
        """POST /redact rejects non-docx file uploads."""
        fake_file = io.BytesIO(b"Hello world text file")
        response = client.post(
            "/redact",
            files={"file": ("sample.txt", fake_file, "text/plain")},
        )
        assert response.status_code == 400
        assert "Invalid file format" in response.json().get("detail", "")

    def test_redact_synthetic_docx(self, client):
        """POST /redact processes synthetic DOCX and removes original PII values."""
        original_name = "Rajesh Kumar"
        original_email = "rajesh.kumar@example.com"
        original_phone = "+91 9876543210"

        doc_io = _create_synthetic_docx([
            f"Employee {original_name} submitted the quarterly audit report.",
            f"Contact the team at {original_email} for follow-ups.",
            f"Primary contact number: {original_phone}.",
        ])

        response = client.post(
            "/redact",
            files={
                "file": (
                    "test_audit.docx",
                    doc_io.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        assert response.status_code == 200
        assert "application/vnd.openxmlformats-officedocument" in response.headers.get("content-type", "")

        # Parse returned DOCX bytes to verify redactions
        out_doc = docx.Document(io.BytesIO(response.content))
        out_text = "\n".join(p.text for p in out_doc.paragraphs)

        # Original PII values must be completely absent from returned text
        assert original_name not in out_text, f"Found '{original_name}' in output text"
        assert original_email not in out_text, f"Found '{original_email}' in output text"
        assert "9876543210" not in out_text, f"Found '{original_phone}' in output text"
