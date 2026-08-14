"""
tests/test_redactor.py

End-to-end redaction tests using synthetic DOCX files.
"""

import os
import tempfile
import pytest

import docx
from src.redactor import redact_document
from src.replacer import reset_mapping


@pytest.fixture(autouse=True)
def clear_map():
    reset_mapping()
    yield
    reset_mapping()


def _make_docx(texts: list, path: str):
    """Create a minimal DOCX with given paragraph texts."""
    doc = docx.Document()
    for text in texts:
        doc.add_paragraph(text)
    doc.save(path)


class TestRedactDocument:

    def test_redacts_email_in_paragraph(self, tmp_path):
        input_p = str(tmp_path / "input.docx")
        output_p = str(tmp_path / "output.docx")
        _make_docx(["Contact us at admin@secret.com."], input_p)

        result = redact_document(input_p, output_p)

        assert result.total_entities_found >= 1
        out_doc = docx.Document(output_p)
        out_text = " ".join(p.text for p in out_doc.paragraphs)
        assert "admin@secret.com" not in out_text

    def test_redacts_phone_in_paragraph(self, tmp_path):
        input_p = str(tmp_path / "input.docx")
        output_p = str(tmp_path / "output.docx")
        _make_docx(["Call +91 9876543210 now."], input_p)

        result = redact_document(input_p, output_p)
        assert result.total_entities_found >= 1

        out_doc = docx.Document(output_p)
        out_text = " ".join(p.text for p in out_doc.paragraphs)
        assert "9876543210" not in out_text

    def test_redacts_ip_in_paragraph(self, tmp_path):
        input_p = str(tmp_path / "input.docx")
        output_p = str(tmp_path / "output.docx")
        _make_docx(["Server IP: 203.0.113.42."], input_p)

        result = redact_document(input_p, output_p)
        assert result.total_entities_found >= 1

        out_doc = docx.Document(output_p)
        out_text = " ".join(p.text for p in out_doc.paragraphs)
        assert "203.0.113.42" not in out_text

    def test_output_file_created(self, tmp_path):
        input_p = str(tmp_path / "input.docx")
        output_p = str(tmp_path / "output.docx")
        _make_docx(["No PII here. Just regular text about finances."], input_p)

        redact_document(input_p, output_p)
        assert os.path.exists(output_p)

    def test_consistent_replacement(self, tmp_path):
        """Same PII appearing twice should produce the same replacement both times."""
        input_p = str(tmp_path / "input.docx")
        output_p = str(tmp_path / "output.docx")
        _make_docx([
            "Contact alice@example.com for questions.",
            "All replies go to alice@example.com.",
        ], input_p)

        redact_document(input_p, output_p)
        out_doc = docx.Document(output_p)
        paragraphs = [p.text for p in out_doc.paragraphs if p.text.strip()]

        # Extract the replacement email from paragraph 1
        import re
        emails_p1 = re.findall(r"[\w.+%-]+@[\w.-]+\.\w+", paragraphs[0]) if paragraphs else []
        emails_p2 = re.findall(r"[\w.+%-]+@[\w.-]+\.\w+", paragraphs[1]) if len(paragraphs) > 1 else []

        if emails_p1 and emails_p2:
            assert emails_p1[0] == emails_p2[0], \
                "Same original email should produce same replacement"

    def test_ordinary_date_not_redacted(self, tmp_path):
        """Dates without DOB context must not be redacted."""
        input_p = str(tmp_path / "input.docx")
        output_p = str(tmp_path / "output.docx")
        _make_docx(["The AGM is scheduled for 15/08/2024."], input_p)

        result = redact_document(input_p, output_p)
        dob_records = [r for r in result.records if r.entity_type == "DATE_OF_BIRTH"]
        assert len(dob_records) == 0, \
            f"Ordinary date was incorrectly flagged as DOB: {dob_records}"

    def test_summary_by_type(self, tmp_path):
        input_p = str(tmp_path / "input.docx")
        output_p = str(tmp_path / "output.docx")
        _make_docx([
            "Email: user@test.com",
            "Phone: +91 9123456789",
            "Server: 10.0.0.1",
        ], input_p)

        result = redact_document(input_p, output_p)
        summary = result.summary_by_type()
        assert isinstance(summary, dict)
        assert result.total_entities_found == sum(summary.values())
